#!/usr/bin/env python3
"""
WikiAccess - WCAG 2.1 AA/AAA Accessibility Checker

Validates HTML and DOCX files for accessibility compliance:
- Heading hierarchy
- Alt text for images
- Color contrast
- Link accessibility
- Document structure

Part of WikiAccess: Transform DokuWiki into Accessible Documents
https://github.com/yourusername/wikiaccess
"""

from pathlib import Path
from typing import Dict, List
from bs4 import BeautifulSoup
import re
from docx import Document


class AccessibilityChecker:
    """Check WCAG 2.1 compliance for HTML and DOCX"""
    
    def __init__(self):
        self.issues_aa = []
        self.issues_aaa = []
        self.warnings = []
        self.passes = []
    
    def check_html(self, html_path: str) -> Dict:
        """
        Check HTML file for accessibility issues
        
        Returns dict with:
            - score_aa: 0-100
            - score_aaa: 0-100
            - issues: list of issues
            - warnings: list of warnings
            - passes: list of passed checks
        """
        self.issues_aa = []
        self.issues_aaa = []
        self.warnings = []
        self.passes = []
        
        with open(html_path, 'r', encoding='utf-8') as f:
            html = f.read()
        
        soup = BeautifulSoup(html, 'html.parser')
        
        # Run checks
        self._check_html_lang(soup)
        self._check_page_title(soup)
        self._check_heading_hierarchy(soup)
        self._check_images(soup)
        self._check_links(soup)
        self._check_color_contrast(soup)
        self._check_semantic_html(soup)
        self._check_skip_links(soup)
        self._check_aria_labels(soup)
        self._check_form_labels(soup)
        self._check_keyboard_access(soup)
        
        # Calculate scores
        total_checks_aa = len(self.issues_aa) + len(self.passes)
        total_checks_aaa = len(self.issues_aaa) + len(self.passes)
        
        score_aa = 100 if total_checks_aa == 0 else int((len(self.passes) / total_checks_aa) * 100)
        score_aaa = 100 if total_checks_aaa == 0 else int((len(self.passes) / (total_checks_aa + total_checks_aaa)) * 100)
        
        return {
            'score_aa': score_aa,
            'score_aaa': score_aaa,
            'issues_aa': self.issues_aa,
            'issues_aaa': self.issues_aaa,
            'warnings': self.warnings,
            'passes': self.passes,
            'file': html_path
        }
    
    def check_docx(self, docx_path: str) -> Dict:
        """Check DOCX file for accessibility"""
        self.issues_aa = []
        self.issues_aaa = []
        self.warnings = []
        self.passes = []
        
        try:
            doc = Document(docx_path)
        except Exception as e:
            return {
                'score_aa': 0,
                'score_aaa': 0,
                'issues_aa': [f"Could not open document: {e}"],
                'issues_aaa': [],
                'warnings': [],
                'passes': [],
                'file': docx_path
            }
        
        # Check document properties
        self._check_docx_title(doc)
        self._check_docx_language(doc)
        self._check_docx_headings(doc)
        self._check_docx_images(doc)
        self._check_docx_links(doc)
        
        # Calculate scores
        total_checks_aa = len(self.issues_aa) + len(self.passes)
        total_checks_aaa = len(self.issues_aaa) + len(self.passes)
        
        score_aa = 100 if total_checks_aa == 0 else int((len(self.passes) / total_checks_aa) * 100)
        score_aaa = 100 if total_checks_aaa == 0 else int((len(self.passes) / (total_checks_aa + total_checks_aaa)) * 100)
        
        return {
            'score_aa': score_aa,
            'score_aaa': score_aaa,
            'issues_aa': self.issues_aa,
            'issues_aaa': self.issues_aaa,
            'warnings': self.warnings,
            'passes': self.passes,
            'file': docx_path
        }
    
    # HTML Checks
    
    def _check_html_lang(self, soup):
        """Check for lang attribute on html element"""
        html_tag = soup.find('html')
        if html_tag and html_tag.get('lang'):
            self.passes.append('✓ HTML lang attribute present')
        else:
            self.issues_aa.append('✗ Missing lang attribute on <html> tag (WCAG 3.1.1)')
    
    def _check_page_title(self, soup):
        """Check for page title"""
        title = soup.find('title')
        if title and title.string and len(title.string.strip()) > 0:
            self.passes.append(f'✓ Page title present: "{title.string}"')
        else:
            self.issues_aa.append('✗ Missing or empty <title> element (WCAG 2.4.2)')
    
    def _check_heading_hierarchy(self, soup):
        """Check heading hierarchy"""
        headings = soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6'])
        
        if not headings:
            self.warnings.append('⚠ No headings found')
            return
        
        # Check for h1
        h1_count = len(soup.find_all('h1'))
        if h1_count == 0:
            self.issues_aa.append('✗ Missing <h1> heading (WCAG 2.4.6)')
        elif h1_count > 1:
            self.warnings.append(f'⚠ Multiple <h1> headings ({h1_count}) - consider using only one')
        else:
            self.passes.append('✓ Single <h1> heading present')
        
        # Check hierarchy
        prev_level = 0
        hierarchy_ok = True
        for heading in headings:
            level = int(heading.name[1])
            if prev_level > 0 and level > prev_level + 1:
                self.issues_aa.append(f'✗ Heading hierarchy skip: {heading.name} after h{prev_level} (WCAG 1.3.1)')
                hierarchy_ok = False
            prev_level = level
        
        if hierarchy_ok and len(headings) > 1:
            self.passes.append(f'✓ Heading hierarchy correct ({len(headings)} headings)')
    
    def _check_images(self, soup):
        """Check images for alt text and include figure names when available"""
        images = soup.find_all('img')
        
        if not images:
            self.passes.append('✓ No images to check')
            return
        
        from pathlib import Path as _Path
        
        for img in images:
            src = img.get('src', 'unknown')
            alt = img.get('alt')
            figure_name = None
            # Try to use figcaption text if present
            if img.parent and getattr(img.parent, 'name', '') == 'figure':
                cap = img.parent.find('figcaption')
                if cap and cap.get_text(strip=True):
                    figure_name = cap.get_text(strip=True)
            # Fallback to filename
            if not figure_name and src and src != 'unknown':
                figure_name = _Path(src).name
            
            if alt is None:
                label = figure_name or src or 'image'
                self.issues_aa.append(f'✗ Image missing alt attribute: {label} (src: {src}) (WCAG 1.1.1)')
            elif len(alt.strip()) == 0:
                label = figure_name or src or 'image'
                self.warnings.append(f'⚠ Image with empty alt text: {label} (src: {src}) - verify if decorative')
        
        if all(img.has_attr('alt') and img.get('alt', '').strip() for img in images):
            self.passes.append(f'✓ All {len(images)} images have descriptive alt text')
    
    def _check_links(self, soup):
        """Check links for accessibility"""
        links = soup.find_all('a')
        
        if not links:
            self.passes.append('✓ No links to check')
            return
        
        issues = []
        
        for link in links:
            text = link.get_text().strip().lower()
            href = link.get('href', '')
            
            # Check for "click here" or "here" links
            if text in ['click here', 'here', 'read more', 'more', 'link']:
                issues.append(f'✗ Non-descriptive link text: "{text}" ({href}) (WCAG 2.4.4)')
            
            # Check for empty links
            if not text:
                issues.append(f'✗ Empty link text: {href} (WCAG 2.4.4)')
        
        if issues:
            self.issues_aa.extend(issues[:5])  # Limit to first 5
            if len(issues) > 5:
                self.warnings.append(f'⚠ {len(issues) - 5} more link issues found')
        else:
            self.passes.append(f'✓ All {len(links)} links have descriptive text')
    
    def _check_color_contrast(self, soup):
        """Check for color contrast issues"""
        # This is a simplified check - full contrast checking requires rendering
        style_tags = soup.find_all('style')
        inline_styles = soup.find_all(style=True)
        
        # Check for low contrast warnings in styles
        low_contrast_patterns = [
            r'color:\s*#[cdefCDEF]{6}',  # Very light colors
            r'color:\s*#[0-3]{6}',  # Very dark colors with potentially light bg
        ]
        
        has_contrast_concern = False
        for style_tag in style_tags:
            for pattern in low_contrast_patterns:
                if re.search(pattern, style_tag.string or ''):
                    has_contrast_concern = True
        
        if has_contrast_concern:
            self.warnings.append('⚠ Potential color contrast issues detected - manual review recommended')
        else:
            self.passes.append('✓ No obvious color contrast issues in CSS')
    
    def _check_semantic_html(self, soup):
        """Check for semantic HTML5 elements"""
        semantic_elements = ['main', 'nav', 'header', 'footer', 'article', 'section']
        found = [elem for elem in semantic_elements if soup.find(elem)]
        
        if 'main' in found:
            self.passes.append('✓ <main> landmark present')
        else:
            self.issues_aa.append('✗ Missing <main> landmark element (WCAG 1.3.1)')
        
        if len(found) >= 2:
            self.passes.append(f'✓ Semantic HTML elements used: {", ".join(found)}')
        else:
            self.warnings.append('⚠ Consider using more semantic HTML5 elements')
    
    def _check_skip_links(self, soup):
        """Check for skip navigation links"""
        skip_link = soup.find('a', class_='skip-link') or soup.find('a', href='#main-content')
        
        if skip_link:
            self.passes.append('✓ Skip navigation link present')
        else:
            self.issues_aaa.append('✗ Missing skip navigation link (WCAG AAA 2.4.1)')
    
    def _check_aria_labels(self, soup):
        """Check ARIA labels"""
        aria_elements = soup.find_all(attrs={'aria-label': True})
        
        if aria_elements:
            self.passes.append(f'✓ {len(aria_elements)} ARIA labels found')
    
    def _check_form_labels(self, soup):
        """Check form labels"""
        inputs = soup.find_all(['input', 'select', 'textarea'])
        
        if not inputs:
            return
        
        unlabeled = []
        for inp in inputs:
            inp_id = inp.get('id')
            has_label = soup.find('label', attrs={'for': inp_id}) if inp_id else None
            has_aria_label = inp.has_attr('aria-label') or inp.has_attr('aria-labelledby')
            
            if not has_label and not has_aria_label:
                unlabeled.append(inp.get('name', inp.get('type', 'unknown')))
        
        if unlabeled:
            self.issues_aa.append(f'✗ {len(unlabeled)} form inputs without labels (WCAG 3.3.2)')
        elif inputs:
            self.passes.append(f'✓ All {len(inputs)} form inputs have labels')
    
    def _check_keyboard_access(self, soup):
        """Check for keyboard accessibility indicators"""
        # Check for focus styles
        style_tags = soup.find_all('style')
        has_focus_styles = False
        
        for style_tag in style_tags:
            if ':focus' in (style_tag.string or ''):
                has_focus_styles = True
                break
        
        if has_focus_styles:
            self.passes.append('✓ Focus styles defined in CSS')
        else:
            self.warnings.append('⚠ No :focus styles found - keyboard navigation may be unclear')
    
    # DOCX Checks
    
    def _check_docx_title(self, doc):
        """Check document title"""
        title = doc.core_properties.title
        if title:
            self.passes.append(f'✓ Document title set: "{title}"')
        else:
            self.issues_aa.append('✗ Missing document title (WCAG 2.4.2)')
    
    def _check_docx_language(self, doc):
        """Check document language"""
        lang = doc.core_properties.language
        if lang:
            self.passes.append(f'✓ Document language set: {lang}')
        else:
            self.issues_aa.append('✗ Missing document language setting (WCAG 3.1.1)')
    
    def _check_docx_headings(self, doc):
        """Check heading styles"""
        headings = [p for p in doc.paragraphs if p.style.name.startswith('Heading')]
        
        if not headings:
            self.issues_aa.append('✗ No heading styles used (WCAG 1.3.1)')
            return
        
        # Check for Heading 1
        h1_count = len([h for h in headings if h.style.name == 'Heading 1'])
        if h1_count == 0:
            self.issues_aa.append('✗ No Heading 1 style used (WCAG 2.4.6)')
        elif h1_count > 1:
            self.warnings.append(f'⚠ Multiple Heading 1 styles ({h1_count})')
        else:
            self.passes.append('✓ Single Heading 1 present')
        
        self.passes.append(f'✓ {len(headings)} heading styles used')
    
    def _check_docx_images(self, doc):
        """Check images for alt text and report figure names when available"""
        # Rough count of images via relationships
        image_count = sum(1 for rel in doc.part.rels.values() if "image" in getattr(rel, 'target_ref', ''))
        if image_count == 0:
            self.passes.append('✓ No images to check')
        
        missing_alt_images = []
        empty_alt_images = []
        good_images = 0
        
        from docx.oxml.ns import qn
        from pathlib import Path as _P
        
        # Inspect inline shapes for alt properties
        for shape in getattr(doc, 'inline_shapes', []):
            try:
                docPr = shape._inline.docPr
                # Prefer explicit properties
                descr = docPr.get('descr') or docPr.get('{http://schemas.openxmlformats.org/drawingml/2006/main}descr')
                title = docPr.get('title')
                name_attr = docPr.get('name')
                
                # Try to resolve actual image filename via relationship
                img_filename = None
                try:
                    blip = shape._inline.graphic.graphicData.pic.blipFill.blip
                    rId = blip.get(qn('r:embed'))
                    if rId and rId in doc.part.related_parts:
                        part = doc.part.related_parts[rId]
                        img_filename = _P(str(part.partname)).name
                except Exception:
                    pass
                
                image_label = name_attr or img_filename or 'Image'
                text = (descr or title)
                if text is None:
                    missing_alt_images.append(image_label)
                elif len(text.strip()) == 0:
                    empty_alt_images.append(image_label)
                else:
                    good_images += 1
            except Exception:
                missing_alt_images.append('Image')
        
        # Report findings
        for img_name in missing_alt_images:
            self.issues_aa.append(f'✗ Image missing alt text: {img_name} (WCAG 1.1.1)')
        for img_name in empty_alt_images:
            self.warnings.append(f'⚠ Image with empty alt text: {img_name} - verify if decorative')
        if good_images > 0:
            self.passes.append(f'✓ {good_images} images have alt text')
    
    def _check_docx_links(self, doc):
        """Check hyperlinks"""
        # Count hyperlinks
        hyperlink_count = 0
        for paragraph in doc.paragraphs:
            if paragraph._element.xpath('.//w:hyperlink'):
                hyperlink_count += 1
        
        if hyperlink_count > 0:
            self.passes.append(f'✓ {hyperlink_count} hyperlinks found')
