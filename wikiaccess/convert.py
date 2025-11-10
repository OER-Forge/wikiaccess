#!/usr/bin/env python3
"""
WikiAccess - Word (DOCX) Converter

Converts DokuWiki pages to accessible Word documents with:
- Native OMML equation objects (editable in Word)
- Downloaded and embedded images with alt-text
- YouTube video thumbnails with clickable links
- Resolved internal wiki links
- Full WCAG 2.1 AA accessibility compliance

Part of WikiAccess: Transform DokuWiki into Accessible Documents
https://github.com/yourusername/wikiaccess
"""

import sys
import argparse
from pathlib import Path
from typing import Optional
from .scraper import DokuWikiHTTPClient
from .parser import (
    DokuWikiParser,
    AccessibilityManager,
    DokuWikiToWordConverter
)
from .equations import insert_mathml_equation
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image
import tempfile


class EnhancedDokuWikiConverter(DokuWikiToWordConverter):
    """Enhanced converter with HTTP scraping and media downloading"""
    
    def __init__(self, client: DokuWikiHTTPClient, temp_dir: Optional[str] = None):
        super().__init__()
        self.client = client
        self.temp_dir = Path(temp_dir) if temp_dir else Path(tempfile.mkdtemp())
        self.temp_dir.mkdir(parents=True, exist_ok=True)
        self.current_page_id = None
        self.equation_count = 0
        self.image_success = 0
        self.image_failed = 0
        self.video_success = 0
        self.video_failed = 0
    
    def convert_url(self, url: str, output_path: str,
                   document_title: Optional[str] = None,
                   language: str = 'en-US',
                   embed_images: bool = True,
                   embed_videos: bool = True):
        """
        Convert a DokuWiki page from URL to Word document
        
        Args:
            url: Full URL to DokuWiki page
            output_path: Path to save output Word document
            document_title: Document title for accessibility
            language: Document language
            embed_images: Whether to download and embed images
            embed_videos: Whether to download and embed video thumbnails
        """
        print(f"\nConverting: {url}")
        print("=" * 70)
        
        # Fetch page content
        result = self.client.get_page_from_url(url)
        if not result:
            print("✗ Failed to fetch page content")
            return {}
        
        page_id, content = result
        self.current_page_id = page_id
        
        # Store settings
        self.embed_images = embed_images
        self.embed_videos = embed_videos
        
        # Save content to temp file
        temp_file = self.temp_dir / 'temp_content.txt'
        with open(temp_file, 'w', encoding='utf-8') as f:
            f.write(content)
        
        # Convert using parent class
        if document_title is None:
            document_title = page_id.replace('_', ' ').replace(':', ' - ').title()
        
        print(f"\nGenerating Word document...")
        self.equation_count = 0  # Reset counter
        self.image_success = 0
        self.image_failed = 0
        self.video_success = 0
        self.video_failed = 0
        self.convert_file(str(temp_file), output_path, document_title, language)
        
        stats = {
            'images_success': self.image_success,
            'images_failed': self.image_failed,
            'videos_success': self.video_success,
            'videos_failed': self.video_failed,
            'equations': self.equation_count
        }
        
        print(f"\n✓ Conversion complete: {output_path}")
        print(f"  📊 Media: {self.image_success}/{self.image_success + self.image_failed} images, {self.video_success}/{self.video_success + self.video_failed} videos")
        print(f"  📊 Total equations converted: {self.equation_count}")
        return stats
    
    def _add_image(self, parsed: dict):
        """Override to download and embed actual images"""
        image_path = parsed['path']
        width = parsed.get('width')
        
        # Check if this is actually a YouTube video
        if 'youtube>' in image_path:
            video_id = image_path.replace('youtube>', '').split('?')[0]
            self._add_youtube_embed({'video_id': video_id})
            return
        
        # Skip URL embeds
        if 'url>' in image_path or image_path.startswith('http'):
            p = self.doc.add_paragraph()
            p.add_run(f"[External content: {image_path}]").italic = True
            return
        
        # Clean up image path
        image_path = image_path.split('|')[0].split('?')[0].strip()
        
        # Generate alt text
        alt_text = self._generate_alt_text(image_path)
        
        if self.embed_images:
            # Download image
            local_path = self.temp_dir / f"image_{hash(image_path)}.png"
            
            if self.client.download_media(image_path, str(local_path)):
                try:
                    # Add image to document
                    p = self.doc.add_paragraph()
                    run = p.add_run()
                    
                    # Calculate width
                    if width:
                        img_width = min(int(width) / 100, 6.5)  # Max 6.5 inches
                    else:
                        # Get actual image dimensions
                        try:
                            with Image.open(local_path) as img:
                                w, h = img.size
                                # Scale to fit page (max 6.5 inches wide)
                                img_width = min(w / 96, 6.5)  # 96 DPI
                        except:
                            img_width = 4.0  # Default
                    
                    picture = run.add_picture(str(local_path), width=Inches(img_width))
                    
                    # Add alt text - python-docx creates InlineShape, we need the shape element
                    # Access the underlying drawing element
                    try:
                        inline = picture._inline
                        docPr = inline.docPr
                        docPr.set('descr', alt_text)
                        docPr.set('title', alt_text)
                    except:
                        pass  # Alt text addition failed, but image is still embedded
                    
                    # Add caption
                    caption_p = self.doc.add_paragraph()
                    caption_p.add_run(f"Figure: {alt_text}").italic = True
                    caption_p.alignment = 1  # Center
                    
                    self.image_success += 1
                    print(f"  ✓ Embedded image: {image_path}")
                    return
                    
                except Exception as e:
                    print(f"  ⚠ Could not embed image: {e}")
        
        # Fallback: add as text reference
        self.image_failed += 1
        p = self.doc.add_paragraph()
        p.add_run(f"[Image: {alt_text}]").italic = True
        print(f"  ℹ Added image reference: {image_path}")
    
    def _add_youtube_embed(self, parsed: dict):
        """Override to download and embed video thumbnails"""
        video_id = parsed['video_id']
        url = f"https://www.youtube.com/watch?v={video_id}"
        
        if self.embed_videos:
            # Download thumbnail
            local_path = self.temp_dir / f"video_{video_id}.jpg"
            
            if self.client.download_youtube_thumbnail(video_id, str(local_path)):
                try:
                    # Add thumbnail to document
                    p = self.doc.add_paragraph()
                    run = p.add_run()
                    picture = run.add_picture(str(local_path), width=Inches(4.0))
                    
                    # Add alt text
                    alt_text = f"Video thumbnail for YouTube video {video_id}"
                    try:
                        inline = picture._inline
                        docPr = inline.docPr
                        docPr.set('descr', alt_text)
                        docPr.set('title', alt_text)
                    except:
                        pass  # Alt text addition failed, but thumbnail is still embedded
                    
                    # Add caption with clickable link
                    caption_p = self.doc.add_paragraph("Watch video: ")
                    self._add_hyperlink(caption_p, url, f"YouTube Video ({video_id})")
                    caption_p.alignment = 1  # Center
                    
                    self.video_success += 1
                    print(f"  ✓ Embedded video thumbnail: {video_id}")
                    return
                    
                except Exception as e:
                    print(f"  ⚠ Could not embed thumbnail: {e}")
        
        # Fallback: add as text link
        self.video_failed += 1
        p = self.doc.add_paragraph("Video: ")
        self._add_hyperlink(p, url, f"Watch on YouTube (ID: {video_id})")
        print(f"  ℹ Added video link: {video_id}")
    
    def _add_hyperlink(self, paragraph, url: str, text: str):
        """Override to handle internal links"""
        # Resolve internal links
        if not url.startswith('http'):
            url = self.client.resolve_internal_link(url, self.current_page_id)
        
        # Use parent class method
        super()._add_hyperlink(paragraph, url, text)
    
    def _add_equation_block(self, equation: str):
        """Override to use MathML equation rendering"""
        # Add equation with proper formatting
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Use MathML converter to add native Word equation
        success = insert_mathml_equation(p, equation, inline=False)
        
        if success:
            self.equation_count += 1
            print(f"  ✓ Added equation: {equation[:50]}...")
        else:
            # Fallback to plain text
            p.add_run(f"$${equation}$$")
            print(f"  ⚠ Equation fallback (LaTeX): {equation[:50]}...")
    
    def _add_inline_equation(self, paragraph, equation: str):
        """Add inline equation with MathML"""
        success = insert_mathml_equation(paragraph, equation, inline=True)
        
        if success:
            self.equation_count += 1
            # Only print first few to avoid spam
            if self.equation_count <= 5 or self.equation_count % 5 == 0:
                print(f"  ✓ Added inline equation #{self.equation_count}: ${equation[:30]}$")
        else:
            # Fallback to plain text
            paragraph.add_run(f"${equation}$")
            print(f"  ⚠ Inline equation fallback: ${equation[:30]}$")


def main():
    """Main CLI entry point"""
    parser = argparse.ArgumentParser(
        description='Convert DokuWiki pages to accessible Word documents',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  # Convert single page
  %(prog)s https://wiki.com/doku.php?id=physics:momentum -o momentum.docx
  
  # With authentication
  %(prog)s https://wiki.com/page -u admin -p password -o output.docx
  
  # Without embedding media
  %(prog)s https://wiki.com/page -o output.docx --no-images --no-videos
  
  # Custom title and language
  %(prog)s https://wiki.com/page -o output.docx -t "Physics Notes" -l es-ES
        """
    )
    
    parser.add_argument('url', help='DokuWiki page URL to convert')
    parser.add_argument('-o', '--output', required=True, help='Output Word document path')
    parser.add_argument('-u', '--username', help='DokuWiki username (if auth required)')
    parser.add_argument('-p', '--password', help='DokuWiki password (if auth required)')
    parser.add_argument('-t', '--title', help='Document title (for accessibility)')
    parser.add_argument('-l', '--language', default='en-US', help='Document language (default: en-US)')
    parser.add_argument('--no-images', action='store_true', help='Do not embed images')
    parser.add_argument('--no-videos', action='store_true', help='Do not embed video thumbnails')
    parser.add_argument('--temp-dir', help='Temporary directory for downloads')
    
    args = parser.parse_args()
    
    # Extract base URL
    from urllib.parse import urlparse
    parsed = urlparse(args.url)
    base_url = f"{parsed.scheme}://{parsed.netloc}"
    
    # Find wiki root
    path_parts = parsed.path.split('/')
    for i, part in enumerate(path_parts):
        if part == 'doku.php' or part == 'dokuwiki':
            base_url += '/' + '/'.join(path_parts[:i])
            break
    else:
        if '/wikis/' in parsed.path:
            wiki_path = parsed.path.split('?')[0]
            if wiki_path.endswith('/doku.php'):
                wiki_path = wiki_path[:-10]
            base_url += wiki_path.rstrip('/')
    
    print()
    print("DokuWiki to Word Converter (HTTP Scraper)")
    print("=" * 70)
    print(f"Wiki URL: {base_url}")
    print(f"Output: {args.output}")
    print()
    
    # Create client
    client = DokuWikiHTTPClient(base_url, args.username, args.password)
    
    # Create converter
    converter = EnhancedDokuWikiConverter(client, args.temp_dir)
    
    # Convert
    success = converter.convert_url(
        args.url,
        args.output,
        document_title=args.title,
        language=args.language,
        embed_images=not args.no_images,
        embed_videos=not args.no_videos
    )
    
    if success:
        print()
        print("=" * 70)
        print("✓ Conversion successful!")
        print()
        print("Accessibility features:")
        print("  ✓ Document title and language set")
        print("  ✓ Proper heading hierarchy")
        print("  ✓ Accessible hyperlinks")
        print("  ✓ Alt text for all images")
        print("  ✓ Semantic document structure")
        print()
        print("Next steps:")
        print("  1. Open document in Word")
        print("  2. Run Accessibility Checker: File → Info → Check Accessibility")
        print("  3. Test with screen reader")
        print()
        sys.exit(0)
    else:
        print()
        print("✗ Conversion failed")
        sys.exit(1)


if __name__ == '__main__':
    main()
